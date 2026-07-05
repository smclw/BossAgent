from __future__ import annotations

from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".txt"}
MAX_EXTRACTED_CHARS = 50000


def is_supported_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _trim(_extract_pdf(path))
    if suffix == ".docx":
        return _trim(_extract_docx(path))
    if suffix in {".xlsx", ".xls"}:
        return _trim(_extract_excel(path))
    if suffix == ".txt":
        return _trim(path.read_text(encoding="utf-8", errors="ignore"))
    raise ValueError(f"Unsupported file type: {suffix}")


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        pages.append(f"\n\n--- PDF Page {index} ---\n{page.extract_text() or ''}")
    return "".join(pages)


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    table_rows = []
    for table in doc.tables:
        for row in table.rows:
            table_rows.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(paragraphs + table_rows)


def _extract_excel(path: Path) -> str:
    sheets = pd.read_excel(path, sheet_name=None)
    chunks = []
    for sheet_name, df in sheets.items():
        chunks.append(f"\n\n--- Sheet: {sheet_name} ---\n")
        chunks.append(df.head(200).to_markdown(index=False))
    return "\n".join(chunks)


def combine_texts(texts: Iterable[str]) -> str:
    return _trim("\n\n".join(texts))


def _trim(text: str) -> str:
    text = text.strip()
    if len(text) <= MAX_EXTRACTED_CHARS:
        return text
    return text[:MAX_EXTRACTED_CHARS] + "\n\n[内容过长，已截断用于本地 MVP 演示。]"
