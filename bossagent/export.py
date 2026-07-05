from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document

from .config import EXPORT_DIR


def save_markdown(report: str, prefix: str = "bossagent-report") -> Path:
    EXPORT_DIR.mkdir(exist_ok=True)
    filename = f"{prefix}-{_stamp()}.md"
    path = EXPORT_DIR / filename
    path.write_text(report, encoding="utf-8")
    return path


def save_docx(report: str, prefix: str = "bossagent-report") -> Path:
    EXPORT_DIR.mkdir(exist_ok=True)
    filename = f"{prefix}-{_stamp()}.docx"
    path = EXPORT_DIR / filename

    doc = Document()
    doc.add_heading("BossAgent 任务报告", level=1)
    for line in report.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped[:3].replace(".", "").isdigit():
            doc.add_paragraph(stripped, style="List Number")
        else:
            doc.add_paragraph(stripped)
    doc.save(path)
    return path


def _stamp() -> str:
    return datetime.now().strftime("%Y%m%d-%H%M%S")
